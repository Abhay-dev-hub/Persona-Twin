"""
Extract raw text from local files: PDF, DOCX, TXT, MD.
"""

from pathlib import Path

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


def extract_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(pages)


def extract_docx(path: Path) -> str:
    import docx  # python-docx

    document = docx.Document(str(path))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

    # Also pull text out of tables, which python-docx's .paragraphs skips.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    return "\n\n".join(paragraphs)


def extract_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def extract_file(path: str | Path) -> str:
    """Dispatch to the right extractor based on file extension."""
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"No such file: {path}")

    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_pdf(path)
    elif ext == ".docx":
        return extract_docx(path)
    elif ext in (".txt", ".md"):
        return extract_txt(path)
    else:
        raise ValueError(
            f"Unsupported file type '{ext}'. Supported: {sorted(SUPPORTED_EXTENSIONS)}"
        )
